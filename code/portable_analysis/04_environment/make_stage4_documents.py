#!/usr/bin/env python3
"""Create the frozen Stage 4 Word deliverables with a consistent decision-memo style."""

from __future__ import annotations

import os

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(os.environ.get("SEPSIS_SIGNATURE_ANALYSIS_ROOT", Path.cwd())).resolve()
NAVY = "17365D"
TEAL = "1F6E6B"
LIGHT = "EAF2F8"
PALE = "EEF6F4"
GRAY = "5B6573"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure(doc: Document, subtitle: str) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in (("Title", 22, NAVY), ("Heading 1", 14, NAVY), ("Heading 2", 11, TEAL)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Stage 4  |  {subtitle}  |  v1.0  |  2026-07-16")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def add_title(doc: Document, title: str, subtitle: str, status: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(title)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run(subtitle)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    box = doc.add_table(rows=1, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.autofit = False
    box.columns[0].width = Cm(3.0)
    box.columns[1].width = Cm(13.8)
    labels = ("冻结状态", status)
    for idx, value in enumerate(labels):
        c = box.cell(0, idx)
        set_cell_shading(c, NAVY if idx == 0 else LIGHT)
        set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = c.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) if idx == 0 else RGBColor.from_string(NAVY)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for i, values in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(values):
            if i % 2 == 1:
                set_cell_shading(cells[j], PALE)
            set_cell_margins(cells[j])
            cells[j].paragraphs[0].add_run(str(value))
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Cm(width)


def force_cjk_font(doc: Document) -> None:
    """LibreOffice respects direct run fonts more reliably than inherited eastAsia styles."""
    def apply_run(run) -> None:
        run.font.name = "Noto Sans CJK SC"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), "Noto Sans CJK SC")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            apply_run(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_run(run)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    apply_run(run)


def pathway_plan() -> None:
    doc = Document()
    configure(doc, "Pathway analysis plan")
    add_title(
        doc,
        "第四阶段通路分析计划",
        "预设通路纵向锚定｜患者内变化｜跨队列汇总",
        "分析结果查看前冻结；仅在可行性门槛通过后执行",
    )

    doc.add_heading("1. 分析目的与解释边界", level=1)
    doc.add_paragraph(
        "本模块用于检验固定宿主RNA签名的患者内评分变化是否与预设宿主反应通路的同期变化一致。"
        "结果属于E2关联证据：可以表述为“与某类生物过程一致”或“与通路变化耦合”，不得表述为通路驱动签名漂移、因果机制或临床效用。"
    )
    add_bullets(doc, [
        "不以签名表现筛选通路，不开展无假设的大范围通路挖掘。",
        "不改变第三阶段固定签名、患者名单、时间窗、评分方向或主要结论。",
        "不同平台在队列内评分；不先合并表达矩阵，不使用跨队列ComBat。",
    ])

    doc.add_heading("2. 正式启动门槛", level=1)
    add_numbered(doc, [
        "至少4个独立队列保留可恢复的全转录组矩阵及患者级时间配对。",
        "T24或T48至少一个主要时间窗由4个队列支持。",
        "每个主要基因集在每个进入队列中的覆盖率至少70%，且至少覆盖10个基因。",
        "主要评分不出现系统性常数分数、极端缺失或平台特异失败。",
    ])
    doc.add_paragraph("任一核心门槛失败时停止正式通路关联；仅保留可行性审计，不以部分队列结果替代预设多队列分析。")

    doc.add_heading("3. 冻结基因集", level=1)
    add_table(doc, ["层级", "来源与版本", "预设集合"], [
        ["主要", "MSigDB Human Hallmark v2026.1.Hs", "Inflammatory response; TNFα/NF-κB; IFN-α; IFN-γ; Complement; Coagulation; Oxidative phosphorylation; Apoptosis; Allograft rejection"],
        ["有限补充", "Reactome current release，下载于2026-07-16", "Neutrophil degranulation; Cross-presentation of particulate exogenous antigens; TCR signaling"],
        ["禁止", "研究结果导向的自建集合", "不创建、不分析"],
    ], [2.0, 4.7, 10.0])

    doc.add_heading("4. 表达预处理与基因映射", level=1)
    add_table(doc, ["队列类型", "冻结处理"], [
        ["Affymetrix / Illumina", "沿用作者处理后的series matrix；使用官方GPL注释；仅保留无歧义HGNC符号；同一基因多探针算术平均。"],
        ["Agilent GSE236713", "通路模块使用作者75th-percentile归一化、global-median baseline-transformed series matrix；样本内秩对全局平移/尺度变换不敏感。固定签名评分仍使用第三阶段原始数据恢复矩阵。"],
        ["RNA-seq GSE110487", "按GENCODE v25将Ensembl gene ID映射至gene_name；基因计数求和；DESeq2盲法VST（design ~1），不使用时间、结局或签名分数。"],
    ], [4.0, 12.7])

    doc.add_heading("5. 主要与敏感性评分", level=1)
    doc.add_paragraph(
        "主要方法为单样本内部秩的centered singscore：在每个样本的全转录组可测基因中计算百分位秩，"
        "取基因集成员平均秩并减0.5。分数越高表示该基因集整体相对表达越高。"
    )
    add_bullets(doc, [
        "敏感性方法：固定α=0.25的ssGSEA rank-walk enrichment score；所有队列使用同一实现。",
        "缺失基因不填0；低于覆盖门槛的队列×通路组合标记不可计算。",
        "同一患者、同一时间窗多样本处理完全沿用第三阶段冻结规则。",
    ])

    doc.add_heading("6. 纵向终点与队列内分析", level=1)
    doc.add_paragraph(
        "患者内通路变化定义为 ΔP(i,t)=P(i,t)−P(i,T0)，主要比较T0–T24与T0–T48。"
        "每个队列×通路×时间窗报告均值、中位数、患者级bootstrap 95%置信区间及有效配对数。"
        "队列内标准化以同一队列、同一通路基线患者分数标准差为分母；基线SD≤10⁻⁸时组合不可计算。"
    )

    doc.add_heading("7. 跨队列汇总与关联", level=1)
    add_bullets(doc, [
        "通路自身变化：先在各队列独立估计，再用REML随机效应模型汇总，报告95%CI、τ²、I²和95%预测区间。",
        "签名–通路耦合：在患者内用Spearman相关；相关系数作Fisher-z转换后跨队列随机效应汇总。",
        "每个签名同时报告全队列结果与排除其开发/可能重叠队列后的结果。",
        "leave-one-dataset-out用于关键结果稳健性诊断。",
    ])

    doc.add_heading("8. 多重比较", level=1)
    add_table(doc, ["分析族", "校正规则"], [
        ["预设Hallmark通路自身变化", "T24与T48分别构成分析族；Benjamini–Hochberg FDR。"],
        ["预设签名×Hallmark耦合", "每个时间窗内8个签名×9个Hallmark为一个分析族；BH-FDR。"],
        ["Reactome补充", "每个时间窗单独BH-FDR，始终标记补充分析。"],
        ["ssGSEA敏感性", "不作为新发现族；用于验证方向与相对结构。"],
    ], [6.0, 10.7])

    doc.add_heading("9. 可报告与不可报告内容", level=1)
    add_table(doc, ["允许", "禁止"], [
        ["“签名变化与炎症/干扰素/抗原呈递通路的同期变化相关。”", "“该通路导致了签名漂移。”"],
        ["“相关方向在多队列中一致/存在异质性。”", "“证明了细胞内机制。”"],
        ["“预测区间提示结果可能依赖临床场景。”", "“可据此指导重复检测或治疗。”"],
    ], [8.3, 8.3])

    doc.add_heading("10. 版本与冻结记录", level=1)
    doc.add_paragraph(
        "计划版本：v1.0；冻结日期：2026-07-16。主要基因集、评分方法、覆盖门槛、时间窗、标准化、"
        "跨队列汇总及多重比较规则在生成任何通路–签名关联结果前锁定。所有源文件记录SHA-256。"
    )
    force_cjk_font(doc)
    doc.save(ROOT / "04_10_pathway_analysis_plan_v1.0.docx")


if __name__ == "__main__":
    pathway_plan()
