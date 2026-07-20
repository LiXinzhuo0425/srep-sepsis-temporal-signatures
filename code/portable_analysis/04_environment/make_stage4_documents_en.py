#!/usr/bin/env python3
"""Create English Stage 4 Word deliverables for reproducible cross-platform rendering."""

from __future__ import annotations

import os

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(os.environ.get("SEPSIS_SIGNATURE_ANALYSIS_ROOT", Path.cwd())).resolve()
NAVY = "17365D"
BLUE = "2C6E9F"
TEAL = "1F6E6B"
PALE = "EAF2F8"
ALT = "F2F7F6"
GRAY = "64748B"


def shade(cell, fill: str) -> None:
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure(doc: Document, footer_text: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color in (("Title", 23, NAVY), ("Heading 1", 14, BLUE), ("Heading 2", 11, TEAL)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Stage 4  |  {footer_text}  |  v1.0  |  2026-07-16")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def title(doc: Document, heading: str, subheading: str, status: str) -> None:
    doc.add_paragraph(heading, style="Title")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(subheading)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.1)
    table.columns[1].width = Cm(13.4)
    for idx, value in enumerate(("STATUS", status)):
        cell = table.cell(0, idx)
        shade(cell, NAVY if idx == 0 else PALE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) if idx == 0 else RGBColor.from_string(NAVY)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)


def numbered(doc: Document, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(headers):
        cell = t.cell(0, j)
        shade(cell, NAVY)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, value in enumerate(row):
            if i % 2:
                shade(cells[j], ALT)
            cells[j].paragraphs[0].add_run(str(value))
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)


def pathway_plan() -> None:
    doc = Document()
    configure(doc, "Pathway analysis plan")
    title(doc, "Frozen Pathway Analysis Plan", "Longitudinal biological anchoring of fixed host-response RNA signatures", "Frozen before any pathway–signature association result was generated")
    doc.add_heading("Purpose and claim boundary", level=1)
    doc.add_paragraph(
        "This conditional E2 module tests whether within-patient changes in fixed RNA-signature scores are associated with concurrent changes in prespecified host-response pathways. Results may be described as biological concordance or coupling; they cannot establish causal mechanism, within-cell transcriptional regulation, clinical utility, or treatment response."
    )
    bullets(doc, [
        "No signature, patient, cohort, time window or scoring direction is changed from the Stage 3 freeze.",
        "Pathways are not selected by observed signature behaviour; no unrestricted pathway discovery is performed.",
        "Expression matrices are processed within cohort and are never batch-corrected across studies before scoring.",
    ])
    doc.add_heading("Formal start gate", level=1)
    numbered(doc, [
        "At least four independent cohorts retain a recoverable whole-transcriptome matrix and patient-level pairing.",
        "At least one primary window (T24 or T48) is supported by four cohorts.",
        "Every primary gene set reaches at least 70% mapped coverage and at least 10 genes in at least four independent cohorts; noncompatible dataset–pathway combinations are excluded.",
        "Primary pathway scores are non-constant and show no systematic platform-specific failure.",
    ])
    doc.add_heading("Frozen gene sets", level=1)
    table(doc, ["Tier", "Source", "Contents"], [
        ["Prespecified primary", "MSigDB Human Hallmark v2026.1.Hs", "Inflammatory response; TNF/NF-kB; IFN-α; IFN-γ; complement; coagulation; oxidative phosphorylation; apoptosis; allograft rejection"],
        ["Secondary complete collection", "MSigDB Human Hallmark v2026.1.Hs", "All 50 Hallmark sets, reported as a separate multiplicity family"],
        ["Limited supplement", "Reactome current release downloaded 2026-07-16", "Neutrophil degranulation; particulate-antigen cross-presentation; TCR signalling"],
    ], [3.3, 5.0, 8.2])
    doc.add_heading("Expression preprocessing and mapping", level=1)
    table(doc, ["Data type", "Frozen rule"], [
        ["Affymetrix / Illumina", "Author-processed series matrix; official GEO platform annotation; unambiguous symbols only; arithmetic mean across probes mapped to one gene."],
        ["Agilent GSE236713", "Author 75th-percentile normalized, global-median baseline-transformed series matrix for pathway ranks. Fixed-signature scoring remains based on the Stage 3 raw-data rescue matrix."],
        ["RNA-seq GSE110487", "GENCODE v25 Ensembl-to-gene-name mapping; summed counts per gene; DESeq2 blind VST with design ~1; no time, outcome or signature information used."],
    ], [4.0, 12.5])
    doc.add_heading("Scoring and longitudinal endpoints", level=1)
    doc.add_paragraph(
        "The primary method is centered singscore: percentile ranks are computed within each sample across its mapped transcriptome; the mean rank of each gene set is reduced by 0.5. The fixed sensitivity method is a rank-walk ssGSEA implementation with alpha=0.25. Missing genes are never imputed."
    )
    bullets(doc, [
        "Within-patient change: deltaP(i,t) = P(i,t) - P(i,T0), for T24 and T48 separately.",
        "Standardization: divide by the same cohort and pathway baseline-patient SD; SD at or below 1e-8 is non-computable.",
        "Cohort result: paired n, mean, median, patient bootstrap 95% CI, direction proportion and ssGSEA sensitivity result.",
    ])
    doc.add_heading("Cross-cohort synthesis and coupling", level=1)
    bullets(doc, [
        "Pathway changes are synthesized with REML random effects and Hartung–Knapp intervals; tau-squared, I-squared and 95% prediction intervals are reported.",
        "Signature–pathway coupling uses within-cohort Spearman correlation. Fisher-z estimates are meta-analysed with the same random-effects framework.",
        "All-cohort and development-overlap-excluded results are shown. Leave-one-dataset-out results diagnose influential cohorts.",
    ])
    doc.add_heading("Multiplicity and reporting", level=1)
    table(doc, ["Analysis family", "Rule"], [
        ["Nine primary Hallmarks", "T24 and T48 are separate BH-FDR families."],
        ["Complete Hallmark collection", "Separate secondary BH-FDR family at each time window."],
        ["Reactome supplement", "Separate supplementary BH-FDR family at each time window."],
        ["Signature–pathway coupling", "Within each time window and tier; report effect size, heterogeneity and prediction interval, not P value alone."],
    ], [6.2, 10.3])
    doc.add_paragraph("Version v1.0 was frozen on 2026-07-16. Gene-set and annotation source files are protected by SHA-256 records in pathway_analysis_lock.json.")
    doc.save(ROOT / "04_10_pathway_analysis_plan_v1.0.docx")


def clinical_plan() -> None:
    doc = Document()
    configure(doc, "Clinical anchoring plan")
    title(doc, "Clinical Anchoring Decision Plan", "Gate-based handling of longitudinal severity variables", "FORMAL MODULE STOPPED by the prespecified availability gate")
    doc.add_heading("Gate decision", level=1)
    doc.add_paragraph(
        "No deposited serial clinical anchor met the frozen requirement of a common Level 1 definition in at least two independent cohorts and approximately 80 time-matched patients. Therefore no formal signature–clinical-change association is fitted. This decision was made from availability and definition evidence, not from observed associations."
    )
    table(doc, ["Candidate anchor", "Audit result", "Decision"], [
        ["Delta SOFA", "Mentioned for GSE236713 but not deposited at sample level; absent elsewhere", "Cancel formal analysis"],
        ["Shock / vasopressor-state change", "Only baseline septic-shock labels; no serial persistence or resolution field", "Cancel formal analysis"],
        ["Delta lactate; Delta CRP/PCT", "No time-matched sample-level variable in two cohorts", "Cancel formal analysis"],
        ["APACHE II / SAPS II", "Static baseline value or high/low category repeated across RNA time points", "Descriptive availability only"],
        ["Mortality / discharge", "Fixed downstream outcome, not concurrent severity change", "Exploratory context only; no prognostic validation"],
    ], [4.2, 8.5, 3.8])
    doc.add_heading("Analyses not performed", level=1)
    bullets(doc, [
        "No correlation or regression between RNA-score change and SOFA, lactate, CRP, PCT, vasopressor state or ventilation change.",
        "No mortality prediction model, cut-point selection, net reclassification, decision curve or treatment-monitoring claim.",
        "No substitution of baseline APACHE II/SAPS II or fixed mortality labels for a synchronous clinical-change anchor.",
    ])
    doc.add_heading("Residual permitted use", level=1)
    doc.add_paragraph(
        "The paper may report the clinical-variable availability audit and use it to explain why biological interpretation is stronger than clinical anchoring in public longitudinal cohorts. Existing mortality labels may be tabulated as cohort descriptors only. They do not modify the Stage 3 stability classification."
    )
    doc.add_heading("Future activation requirements", level=1)
    numbered(doc, [
        "A single, prespecified serial anchor with harmonized units and sampling windows in at least two independent cohorts.",
        "At least 80 paired patients with RNA and clinical measurements aligned to T0–T24 or T0–T48.",
        "Predeclared covariates, missingness model and multiplicity family before association results are viewed.",
        "Patient-level, cohort-first estimation followed by random-effects synthesis; no pooled record-level model across studies."
    ])
    doc.add_paragraph("The stopped status is recorded in 04_17_signature_clinical_anchor_analysis.xlsx; zero models and zero P values were generated.")
    doc.save(ROOT / "04_16_clinical_anchoring_plan_v1.0.docx")


def prospective_framework() -> None:
    doc = Document()
    configure(doc, "Prospective validation framework")
    title(doc, "Prospective Validation Framework", "Time-locked verification of fixed host-response RNA signatures in suspected sepsis", "Blueprint only; prospective validation not yet performed")
    doc.add_heading("Objective", level=1)
    doc.add_paragraph(
        "Validate analytical repeatability, within-patient T0–T24–T48 score behaviour and concordance with contemporaneous clinical-state change for the eight locked signatures, without re-training coefficients or selecting genes. The principal purpose is interpretation of repeated measurements, not development of a new diagnostic or prognostic model."
    )
    doc.add_heading("Population and sampling", level=1)
    table(doc, ["Element", "Prospective specification"], [
        ["Setting", "Consecutive adults evaluated for suspected sepsis in emergency and ICU pathways; enrol before final adjudication whenever feasible."],
        ["Blood schedule", "T0 within 12 h of recognition/enrolment, T24 at 12–36 h, T48 at 36–60 h. Record exact draw time and treatment exposure."],
        ["Reference standard", "Independent clinical expert panel using predefined infection and organ-dysfunction evidence; adjudicators blinded to RNA score."],
        ["Assay", "Locked RT-qPCR or automated cartridge platform; fixed primer/probe lot rules; predefined normalization and invalid-run criteria."],
        ["Core clinical data", "SOFA components, vasopressor dose, lactate, PCT, CRP, ventilation, antimicrobials, source control, microbiology and discharge/death."],
    ], [4.2, 12.3])
    doc.add_heading("Endpoints", level=1)
    numbered(doc, [
        "Primary: within-patient standardized score change from T0 to T24 and T48 for each fixed signature.",
        "Key clinical anchor: within-patient Delta SOFA matched to the same interval; vasopressor-state change and Delta lactate are secondary anchors.",
        "Analytical repeatability: duplicate-sample coefficient of variation, intraclass correlation and invalid-run proportion.",
        "Interpretation stability: direction consistency, prediction interval, and predeclared movement across score-interpretation bands; no outcome-optimized threshold."
    ])
    doc.add_heading("Bias control and governance", level=1)
    bullets(doc, [
        "Prespecify assay, signature functions, time windows, QC, exclusions, missingness handling and clinical anchor before database lock.",
        "Separate clinical adjudication, laboratory processing and analysis roles; maintain a timestamped deviation log.",
        "Do not impute post-death molecular values. Distinguish death, early discharge, transfer, clinical inability and technical failure.",
        "Report each signature by intended use; do not construct a cross-purpose winner ranking or a new composite stability model."
    ])
    doc.add_heading("Analysis and sample-size planning", level=1)
    doc.add_paragraph(
        "Estimate cohort-level paired changes with patient bootstrap intervals and mixed models, followed by fixed protocol sensitivity analyses. The definitive sample size should use the Stage 3/4 observed paired SD and the smallest clinically interpretable change selected before enrolment. For a two-sided paired comparison, the planning approximation is n = ((z0.975 + z0.80) × SDchange / deltaMCID)^2, inflated for multiplicity, assay invalidity and informative attrition."
    )
    table(doc, ["Planning scenario", "Standardized effect", "Complete pairs before attrition", "With 20% attrition"], [
        ["Moderate change", "0.35 SD", "Approximately 65", "Approximately 82"],
        ["Small-to-moderate change", "0.25 SD", "Approximately 126", "Approximately 158"],
        ["Small change", "0.20 SD", "Approximately 197", "Approximately 247"],
    ], [5.0, 3.2, 4.0, 4.0])
    doc.add_paragraph(
        "A pragmatic multicentre target of 250–300 enrolled patients should be considered if the study must support T48 attrition, assay failures, multiple signatures and clinical-state strata. Final power must be recalculated from the frozen Stage 4 variance estimates and chosen primary signature family."
    )
    doc.add_heading("Decision framework", level=1)
    bullets(doc, [
        "Repeated testing is supported only if analytical repeatability is adequate and the time effect has a sufficiently narrow prediction interval.",
        "Monitoring claims require prespecified, replicated concordance with Delta SOFA or another Level 1 anchor; association alone does not establish treatment responsiveness.",
        "Clinical decision support requires a subsequent impact study and cannot be inferred from this observational validation."
    ])
    doc.save(ROOT / "04_19_prospective_validation_framework.docx")


if __name__ == "__main__":
    pathway_plan()
    clinical_plan()
    prospective_framework()
